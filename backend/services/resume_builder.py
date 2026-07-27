import os
import time
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


class ResumeBuilderService:
    """AI Resume Builder & Optimizer 2.0 supporting PDF, DOCX generation, and LinkedIn summaries."""

    TEMPLATES = {
        "software_engineer": "Software Engineer (Full Stack & Python Specialist)",
        "ai_engineer": "AI & Machine Learning Engineer",
        "data_scientist": "Data Scientist & Analytics Architect",
        "fresher": "Entry-Level Graduate / Fresher",
        "experienced": "Senior Experienced Engineer"
    }

    def __init__(self):
        self.output_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "reports", "interview_reports"))
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_resume_content(self, fullname: str, role_template: str = "ai_engineer", skills: list = None) -> dict:
        """Generates AI optimized resume sections."""
        skills_list = skills or ["Python", "Flask", "SQLAlchemy", "PyTorch", "Docker", "REST APIs"]
        summary = (
            f"Results-driven {self.TEMPLATES.get(role_template, 'Software Engineer')} with expertise in "
            f"{', '.join(skills_list[:4])}. Proven track record designing high-availability web services, "
            "implementing machine learning models, and optimizing database performance."
        )

        experience = [
            {
                "title": f"Lead {role_template.replace('_', ' ').title()} Developer",
                "company": "Tech Corp Solutions",
                "period": "2024 - Present",
                "highlights": [
                    "Architected scalable Flask RESTful microservices handling 50,000 requests/day.",
                    "Integrated computer vision and natural language processing pipelines with 94% precision.",
                    "Optimized SQL database query latency by 45% using indexing and SQLAlchemy ORM pooling."
                ]
            }
        ]

        linkedin_summary = (
            f"Passionate {role_template.replace('_', ' ').title()} building next-generation AI and web applications. "
            f"Specializing in {', '.join(skills_list[:3])}. Open to high-impact software engineering opportunities!"
        )

        return {
            "fullname": fullname,
            "role_template": role_template,
            "summary": summary,
            "skills": skills_list,
            "experience": experience,
            "linkedin_summary": linkedin_summary
        }

    def generate_docx_resume(self, content: dict, user_id: int) -> str:
        """Generates DOCX resume file."""
        doc = Document()
        doc.add_heading(content["fullname"], level=0)
        doc.add_paragraph(content["summary"])

        doc.add_heading("Technical Skills", level=1)
        doc.add_paragraph(", ".join(content["skills"]))

        doc.add_heading("Professional Experience", level=1)
        for exp in content["experience"]:
            doc.add_heading(f"{exp['title']} - {exp['company']} ({exp['period']})", level=2)
            for h in exp["highlights"]:
                doc.add_paragraph(h, style="List Bullet")

        filename = f"resume_{user_id}_{int(time.time())}.docx"
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        return f"reports/interview_reports/{filename}"

    def generate_pdf_resume(self, content: dict, user_id: int) -> str:
        """Generates PDF resume file using ReportLab."""
        filename = f"resume_{user_id}_{int(time.time())}.pdf"
        file_path = os.path.join(self.output_dir, filename)

        c = canvas.Canvas(file_path, pagesize=letter)
        c.setFont("Helvetica-Bold", 18)
        c.drawString(50, 750, content["fullname"])

        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, 725, f"Target Role: {content['role_template'].replace('_', ' ').title()}")

        c.setFont("Helvetica", 10)
        c.drawString(50, 700, "Summary:")
        c.drawString(50, 685, content["summary"][:90])

        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, 650, "Technical Skills:")
        c.setFont("Helvetica", 10)
        c.drawString(50, 635, ", ".join(content["skills"]))

        c.save()
        return f"reports/interview_reports/{filename}"
